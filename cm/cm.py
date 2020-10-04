# -*- coding: utf-8 -*-
"""
Created on Sun Oct  4 01:35:56 2020

@author: chen
"""


import random 
import csv
from docx import Document,styles
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE


##厘米和米比较大小
##包含100厘米，或者多少米多少厘米，米的单位为10以内，厘米单位为100以内。

##带加减法计算，求得数
##带运算比大小


##带直接换算
##带多个比较大小


##设计的规则应该是生成所有可能的数值，然后进行随机取数进行组合
##题型设计：1、加减法计算，以

##随机取出1,200米范围，只有200个
def Mi():
    result = []
    for i in range(1,20) :
        m='%d'%i+'米'
        result.append((m,i))
    return result

MI = Mi()

##随机取出50,200厘米范围，只有150个
def Limi():
    result = []
    for i in range(50,200) :
        m='%d'%i+'厘米'
        result.append((m,i/100))
    return result

LIMI = Limi()

##随机取出50,200厘米范围,加上1-50米以内，有7350个
def Milimi():
    result = []
    for i in range(1,100) :
        for k in range(1,20):
            m='%d'%k+'米''%d'%i+'厘米'
            result.append((m,k+i/100))
    return result

MILIMI = Milimi()


# 加减法，设计随机取前100个，进行计算
def Jia_jian():
    shizi = []
    random.shuffle(MI)
    random.shuffle(LIMI)
    random.shuffle(MILIMI)
 ####切片是冒号，不要搞错了
    MI100=MI   
    LIMI100=LIMI[0:100]
    MILIMI100=MILIMI[0:100]
    
    for i in range(len(MI100)):
        for k in  range(len(LIMI100)):
            if MI100[i][1]-LIMI100[k][1]>0:
                zuhe=str(MI100[i][0])+' - '+str(LIMI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe)
    
    for i in range(len(LIMI100)):
        for k in range(len(MI100)):
            if LIMI100[i][1]-MI100[k][1]>0:
                zuhe2=str(LIMI100[i][0])+' - '+str(MI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe2)

    for i in range(len(MI100)):
        for k in range(len(MILIMI100)):
            if MI100[i][1]-MILIMI100[k][1]>0:
                zuhe3=str(MI100[i][0])+' - '+str(MILIMI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe3)
    
    for i in range(len(MILIMI100)):
        for k in range(len(MI100)):
            if MILIMI100[i][1]-MI100[k][1]>0:
                zuhe4=str(MILIMI100[i][0])+' - '+str(MI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe4)

    for i in range(len(LIMI100)):
        for k in range(len(MILIMI100)):
            if LIMI100[i][1]-MILIMI100[k][1]>0:
                zuhe5=str(LIMI100[i][0])+' - '+str(MILIMI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe5)
    
    for i in range(len(MILIMI100)):
        for k in range(len(LIMI100)):
            if MILIMI100[i][1]-LIMI100[k][1]>0:
                zuhe6=str(MILIMI100[i][0])+' - '+str(LIMI100[k][0])+' = '+'(        )米(        )厘米'
                shizi.append(zuhe6)

    return(shizi)

JIA_JIAN=Jia_jian()


def Yunsuan():
    shizi = []
    random.shuffle(MI)
    random.shuffle(LIMI)
    random.shuffle(MILIMI)
 ####切片是冒号，不要搞错了
    MI100=MI   
    LIMI100=LIMI[0:100]
    MILIMI100=MILIMI[0:100]
    
    for i in range(len(MI100)):
        for k in range(len(LIMI100)):
            for j in range(len(MILIMI100)):
                if MI100[i][1]-LIMI100[k][1]>0:
                    zuhe=str(MI100[i][0])+' - '+str(LIMI100[k][0])+' (    )'+str(MILIMI100[j][0])
                    shizi.append(zuhe)

    for i in range(len(LIMI100)):
        for k in range(len(MI100)):
            for j in range(len(MILIMI100)):
                if LIMI100[i][1]-MI100[k][1]>0:
                    zuhe2=str(LIMI100[i][0])+' - '+str(MI100[k][0])+' (    ) '+str(MILIMI100[j][0])
                    shizi.append(zuhe2)    
                    
    for i in range(len(MI100)):
        for k in range(len(MILIMI100)):
            for j in range(len(LIMI100)):
                if MI100[i][1]-MILIMI100[k][1]>0:
                    zuhe3=str(MI100[i][0])+' - '+str(MILIMI100[k][0])+' (    ) '+str(LIMI100[j][0])
                    shizi.append(zuhe3)

    for i in range(len(MILIMI100)):
        for k in range(len(MI100)):
            for j in range(len(LIMI100)):
                if MILIMI100[i][1]-MI100[k][1]>0:
                    zuhe4=str(MILIMI100[i][0])+' - '+str(MI100[k][0])+' (    ) '+str(LIMI100[j][0])
                    shizi.append(zuhe4)    

    for i in range(len(MILIMI100)):
        for k in range(len(LIMI100)):
            for j in range(len(MI100)):
                if MILIMI100[i][1]-LIMI100[k][1]>0:
                    zuhe5=str(MILIMI100[i][0])+' - '+str(LIMI100[k][0])+' (    ) '+str(MI100[j][0])
                    shizi.append(zuhe5)

    for i in range(len(LIMI100)):
        for k in range(len(MILIMI100)):
            for j in range(len(MI100)):
                if LIMI100[i][1]-MILIMI100[k][1]>0:
                    zuhe6=str(LIMI100[i][0])+' - '+str(MILIMI100[k][0])+' (    ) '+str(MI100[j][0])
                    shizi.append(zuhe6)    

    return(shizi)

YUNSUAN=Yunsuan()


if __name__ == "__main__" :


#分别是两种类型题目，所以分开
    result = []


#随机顺序6种题型，加入到result中来，前面的6种都是全部遍历的，
###，题库只要200题，200题，100题，100题，200题，200题，只要是符合的都在题库当中了
    random.shuffle(JIA_JIAN)
    random.shuffle(YUNSUAN)

    result.extend(JIA_JIAN[0:199])
    result.extend(YUNSUAN[0:199])



###再次随机顺序
    random.shuffle(result)

###生成word的代码
    doc_yy = Document()
    styles = doc_yy.styles
###增加新的模板，字体等
    some1 = styles.add_style('some1',WD_STYLE_TYPE.PARAGRAPH)
    some1.font.size=Pt(14)

####取10题，为一个页面，取result总长度的少10
    i = 0 
    while i < len(result)-20:
        doc_yy.add_paragraph(result[i],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+1],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+2],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+3],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+4],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+5],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+6],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+7],style=some1)
        # doc_yy.add_paragraph(text='\r')   
        doc_yy.add_paragraph(result[i+8],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+9],style=some1)
        doc_yy.add_paragraph()        
        doc_yy.add_paragraph(result[i+10],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+11],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+12],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+13],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+14],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+15],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+16],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+17],style=some1)
        # doc_yy.add_paragraph(text='\r')   
        doc_yy.add_paragraph(result[i+18],style=some1)
        # doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+19],style=some1)

        i = i + 20

    doc_yy.save('CM.docx')




































