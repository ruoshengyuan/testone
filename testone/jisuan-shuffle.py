
import random 
import csv
from docx import Document,styles
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE


def Night():
    result = []
    for i in range(1,10) :
        for m in range(1,10):
            c = i*m
            result.append((i,m,c))
    return result

NIGHT = Night()

def Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][2]<=100 and NIGHT[i][2]>9:
                shizi = str(m)+'+'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan = m+NIGHT[i][2]
                result.append((shizi,daan))
                result.append((shizi2,daan))
    return result

JIA_9 =  Jia_9()

def Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][2]>=0 and NIGHT[i][2]>9:
                shizi = str(m)+'-'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][2]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][2]-m>=0  and NIGHT[i][2]>9:
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][2]-m
                result.append((shizi2,daan2))
    return result

JIAN_9 = Jian_9()


def Jia_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            for n in range (10,100):
                if i+m+n<=100 :
                    shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
                    daan = i+m+n
                    result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i+m<=100  :
                for n in range (10,100):
                    if i+m-n>=0 :
                        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
                        daan3 = i+m-n
                        result.append((shizi3,daan3))
    return result

JIA_QIAN = Jia_qian()

def Jian_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0 :
                for n in range (10,100):
                    if i-m-n>=0 :
                        shizi = str(i)+'-'+str(m)+'-'+str(n)+'='
                        daan = i-m-n
                        result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0  :
                for n in range (10,100):
                    if i-m+n<=100 :
                        shizi3 = str(i)+'-'+str(m)+'+'+str(n)+'='
                        daan3 = i-m+n
                        result.append((shizi3,daan3))
    return result

JIAN_QIAN = Jian_qian()

def Chu_Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][0]>=0 and NIGHT[i][1]>1:
                shizi = str(m)+'-'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]-m>=0 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][0]-m
                result.append((shizi2,daan2))
    return result

CHU_JIAN_9 = Chu_Jian_9()

def Chu_Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][0]<=100 and NIGHT[i][1]>1:
                shizi = str(m)+'+'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m+NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]+m<=100 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan2 = NIGHT[i][0]+m
                result.append((shizi2,daan2))
    return result

CHU_JIA_9 = Chu_Jia_9()

def Jia_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(100,799) 
        m = random.randint(100,899-i)
        n = random.randint(100,1000-m-i)
        shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
        daan = i+m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(100,899) 
        m = random.randint(100,1000-i)
        n = random.randint(100,m+i)
        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
        daan3 = i+m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIA_QIAN1000 = Jia_qian1000()

def Jian_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(201,1000) 
        m = random.randint(100,i)
        n = random.randint(100,1000+m-i)
        shizi = str(i)+'-'+str(m)+'+'+str(n)+'='
        daan = i-m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(301,1000) 
        m = random.randint(100,i-100)
        n = random.randint(100,i-m)
        shizi3 = str(i)+'-'+str(m)+'-'+str(n)+'='
        daan3 = i-m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIAN_QIAN1000 = Jian_qian1000()


if __name__ == "__main__" :


    result = []
    result2 = []

    random.shuffle(JIA_9)
    random.shuffle(JIAN_9)
    random.shuffle(JIA_QIAN)
    random.shuffle(JIAN_QIAN)
    random.shuffle(CHU_JIA_9)
    random.shuffle(CHU_JIAN_9)
    result.extend(JIA_9[0:199])
    result.extend(JIAN_9[0:199])
    result.extend(JIA_QIAN[0:99])
    result.extend(JIAN_QIAN[0:99])
    result.extend(CHU_JIA_9[0:199])
    result.extend(CHU_JIAN_9[0:199])

    random.shuffle(JIA_QIAN1000)
    random.shuffle(JIAN_QIAN1000)
    result2.extend(JIA_QIAN1000[0:499])
    result2.extend(JIAN_QIAN1000[0:499])


    random.shuffle(result)
    random.shuffle(result2)  

    doc_yy = Document()
    styles = doc_yy.styles
    some1 = styles.add_style('some1',WD_STYLE_TYPE.PARAGRAPH)
    some1.font.size=Pt(14)

    i = 0 
    while i < len(result)-10:
        doc_yy.add_paragraph(result[i][0]+' '*50+result2[i][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+1][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+2][0]+' '*50+result2[i+1][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+3][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+4][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+5][0]+' '*50+result2[i+2][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+6][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+7][0],style=some1)
        doc_yy.add_paragraph(text='\r')   
        doc_yy.add_paragraph(result[i+8][0]+' '*50+result2[i+3][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+9][0],style=some1)

        i = i + 10

    doc_yy.save('jieguo.docx')


