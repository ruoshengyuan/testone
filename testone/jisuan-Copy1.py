
import random 
import csv
from docx import Document,styles
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE


def Night():
    result = []
    for i in range(1,10) :
        for m in range(1,10):
            c = i*m
            result.append((i,m,c))
    return result

NIGHT = Night()

def Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][2]<=100 and NIGHT[i][2]>9:
                shizi = str(m)+'+'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan = m+NIGHT[i][2]
                result.append((shizi,daan))
                result.append((shizi2,daan))
    return result

JIA_9 =  Jia_9()

def Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][2]>=0 and NIGHT[i][2]>9:
                shizi = str(m)+'-'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][2]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][2]-m>=0  and NIGHT[i][2]>9:
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][2]-m
                result.append((shizi2,daan2))
    return result

JIAN_9 = Jian_9()


def Jia_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            for n in range (10,100):
                if i+m+n<=100 :
                    shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
                    daan = i+m+n
                    result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i+m<=100  :
                for n in range (10,100):
                    if i+m-n>=0 :
                        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
                        daan3 = i+m-n
                        result.append((shizi3,daan3))
    return result

JIA_QIAN = Jia_qian()

def Jian_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0 :
                for n in range (10,100):
                    if i-m-n>=0 :
                        shizi = str(i)+'-'+str(m)+'-'+str(n)+'='
                        daan = i-m-n
                        result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0  :
                for n in range (10,100):
                    if i-m+n<=100 :
                        shizi3 = str(i)+'-'+str(m)+'+'+str(n)+'='
                        daan3 = i-m+n
                        result.append((shizi3,daan3))
    return result

JIAN_QIAN = Jian_qian()

def Chu_Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][0]>=0 and NIGHT[i][1]>1:
                shizi = str(m)+'-'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]-m>=0 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][0]-m
                result.append((shizi2,daan2))
    return result

CHU_JIAN_9 = Chu_Jian_9()

def Chu_Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][0]<=100 and NIGHT[i][1]>1:
                shizi = str(m)+'+'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m+NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]+m<=100 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan2 = NIGHT[i][0]+m
                result.append((shizi2,daan2))
    return result

CHU_JIA_9 = Chu_Jia_9()

def Jia_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(100,799) 
        m = random.randint(100,899-i)
        n = random.randint(100,1000-m-i)
        shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
        daan = i+m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(100,899) 
        m = random.randint(100,1000-i)
        n = random.randint(100,m+i)
        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
        daan3 = i+m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIA_QIAN1000 = Jia_qian1000()

def Jian_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(201,1000) 
        m = random.randint(100,i)
        n = random.randint(100,1000+m-i)
        shizi = str(i)+'-'+str(m)+'+'+str(n)+'='
        daan = i-m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(301,1000) 
        m = random.randint(100,i-100)
        n = random.randint(100,i-m)
        shizi3 = str(i)+'-'+str(m)+'-'+str(n)+'='
        daan3 = i-m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIAN_QIAN1000 = Jian_qian1000()


if __name__ == "__main__" :
    # print('chengfajiafa',len(JIA_9),
    #         'chengfajianfa',len(JIAN_9),
    #         'chufajiafa',len(CHU_JIA_9),
    #         'chufajianfa',len(CHU_JIAN_9),
    #         'jiaqian',len(JIA_QIAN),
    #         'jianqian',len(JIAN_QIAN)
    #         )

    chengfajiafa = 200
    chengfajianfa = 200
    chufajiafa =200
    chufajianfa = 200
    jiaqian = 100
    jianqian = 100
    jiaqian1000 = 500
    jianqian1000 = 500
    result = []
    result2 = []

    xuhao = []
    i = 0
    while i < chengfajiafa :
        a = random.randint(0,len(JIA_9)-1)
        if a not in xuhao : 
            result.append(JIA_9[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < chengfajianfa :
        a = random.randint(0,len(JIAN_9)-1)
        if a not in xuhao : 
            result.append(JIAN_9[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < chufajiafa :
        a = random.randint(0,len(CHU_JIA_9)-1)
        if a not in xuhao : 
            result.append(CHU_JIA_9[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < chufajianfa :
        a = random.randint(0,len(CHU_JIAN_9)-1)
        if a not in xuhao : 
            result.append(CHU_JIAN_9[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < jiaqian :
        a = random.randint(0,len(JIA_QIAN)-1)
        if a not in xuhao : 
            result.append(JIA_QIAN[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < jianqian :
        a = random.randint(0,len(JIAN_QIAN)-1)
        if a not in xuhao : 
            result.append(JIAN_QIAN[a])
            xuhao.append(a)
            i = i +1

    xuhao = []
    i = 0
    while i < jiaqian1000 :
        a = random.randint(0,len(JIA_QIAN1000)-1)
        if a not in xuhao : 
            result2.append(JIA_QIAN1000[a])
            xuhao.append(a)
            i = i +1
    xuhao = []
    i = 0
    while i < jianqian1000 :
        a = random.randint(0,len(JIAN_QIAN1000)-1)
        if a not in xuhao : 
            result2.append(JIAN_QIAN1000[a])
            xuhao.append(a)
            i = i +1

    result_new2 = []
    result_new = []
    xuhao_new = []
    i = 0
    while i < len(result) :
        b = random.randint(0,len(result)-1)
        if b not in xuhao_new : 
            result_new.append(result[b][0])
            xuhao_new.append(b)
            i = i +1    
    # print(result_new)

    xuhao_new = []
    i = 0
    while i < len(result2) :
        b = random.randint(0,len(result2)-1)
        if b not in xuhao_new : 
            result_new2.append(result2[b][0])
            xuhao_new.append(b)
            i = i +1    
    # print(result_new)
    # with open('shiti.csv','w')  as csvfile:
    #     writer = csv.writer(csvfile) 

    #     writer.writerow(result_new)
    #     # item = 0 
    #     # while item < len(result_new) :
    #     #     writer.writerow(result_new[item])
    #     #     i = i + 1
    #     print("Save,done")
    doc_yy = Document()

    styles = doc_yy.styles
    some1 = styles.add_style('some1',WD_STYLE_TYPE.PARAGRAPH)
    some1.font.size=Pt(14)

    i = 0 
    while i < len(result_new)-10:
        doc_yy.add_paragraph(result_new[i]+' '*50+result_new2[i],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+1],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+2]+' '*50+result_new2[i+1],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+3],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+4],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+5]+' '*50+result_new2[i+2],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+6],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+7],style=some1)
        doc_yy.add_paragraph(text='\r')   
        doc_yy.add_paragraph(result_new[i+8]+' '*50+result_new2[i+3],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result_new[i+9],style=some1)

        i = i + 10

    doc_yy.save('jieguo.docx')


