
import random 
import csv
from docx import Document,styles
from docx.shared import Pt,Inches
from docx.enum.style import WD_STYLE_TYPE


##99乘法试题
def Night():
    result = []
    for i in range(1,10) :
        for m in range(1,10):
            c = i*m
            result.append((i,m,c))
    return result

NIGHT = Night()

##100以内的加法，以99乘法的为其中一个加，
##另外一个m为100以内数字，进行加和，要保证结果仍在100以内
def Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][2]<=100 and NIGHT[i][2]>9:
                shizi = str(m)+'+'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan = m+NIGHT[i][2]
                result.append((shizi,daan))
                result.append((shizi2,daan))
    return result


JIA_9 =  Jia_9()

##100以内的减法，以99乘法的为其中一个加，
##另外一个m为100以内数字，进行减法，要保证结果仍在100以内
def Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][2]>=0 and NIGHT[i][2]>9:
                shizi = str(m)+'-'+str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][2]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][2]-m>=0  and NIGHT[i][2]>9:
                shizi2 = str(NIGHT[i][0])+'x'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][2]-m
                result.append((shizi2,daan2))
    return result

JIAN_9 = Jian_9()

##连加连减，加法在前
##要保证结果仍在100以内
def Jia_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            for n in range (10,100):
                if i+m+n<=100 :
                    shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
                    daan = i+m+n
                    result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i+m<=100  :
                for n in range (10,100):
                    if i+m-n>=0 :
                        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
                        daan3 = i+m-n
                        result.append((shizi3,daan3))
    return result

JIA_QIAN = Jia_qian()


##连加连减，减法在前
##要保证结果仍在100以内
def Jian_qian():
    result = []
    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0 :
                for n in range (10,100):
                    if i-m-n>=0 :
                        shizi = str(i)+'-'+str(m)+'-'+str(n)+'='
                        daan = i-m-n
                        result.append((shizi,daan))

    for i in range (10,100):
        for m in range (10,100):
            if i-m >= 0  :
                for n in range (10,100):
                    if i-m+n<=100 :
                        shizi3 = str(i)+'-'+str(m)+'+'+str(n)+'='
                        daan3 = i-m+n
                        result.append((shizi3,daan3))
    return result

JIAN_QIAN = Jian_qian()



##99乘法表对应的除法运算，包含减法
##要保证结果仍在100以内
def Chu_Jian_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m-NIGHT[i][0]>=0 and NIGHT[i][1]>1:
                shizi = str(m)+'-'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m-NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]-m>=0 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'-'+str(m)+'='
                daan2 = NIGHT[i][0]-m
                result.append((shizi2,daan2))
    return result

CHU_JIAN_9 = Chu_Jian_9()



##99乘法表对应的除法运算，包含加法
##要保证结果仍在100以内
def Chu_Jia_9():
    result = []
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if m+NIGHT[i][0]<=100 and NIGHT[i][1]>1:
                shizi = str(m)+'+'+str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'='
                daan = m+NIGHT[i][0]
                result.append((shizi,daan))
    for i in range(len(NIGHT)) :
        for m in range (1,100):
            if NIGHT[i][0]+m<=100 and NIGHT[i][1]>1:
                shizi2 = str(NIGHT[i][2])+'/'+str(NIGHT[i][1])+'+'+str(m)+'='
                daan2 = NIGHT[i][0]+m
                result.append((shizi2,daan2))
    return result

CHU_JIA_9 = Chu_Jia_9()



##1000以内的加减法
##要保证结果仍在1000以内，这个应该是算成竖式，1000道题目以内
def Jia_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(100,799) 
        m = random.randint(100,899-i)
        n = random.randint(100,1000-m-i)
        shizi = str(i)+'+'+str(m)+'+'+str(n)+'='
        daan = i+m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(100,899) 
        m = random.randint(100,1000-i)
        n = random.randint(100,m+i)
        shizi3 = str(i)+'+'+str(m)+'-'+str(n)+'='
        daan3 = i+m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIA_QIAN1000 = Jia_qian1000()

##1000以内的加减法
##要保证结果仍在1000以内，这个应该是算成竖式，1000道题目以内
def Jian_qian1000():
    result2 = []
    k = 0
    while k < 1000 :
        i = random.randint(201,1000) 
        m = random.randint(100,i)
        n = random.randint(100,1000+m-i)
        shizi = str(i)+'-'+str(m)+'+'+str(n)+'='
        daan = i-m+n
        result2.append((shizi,daan))
        k = k+1
        
    k = 0
    while k < 1000 :
        i = random.randint(301,1000) 
        m = random.randint(100,i-100)
        n = random.randint(100,i-m)
        shizi3 = str(i)+'-'+str(m)+'-'+str(n)+'='
        daan3 = i-m-n
        result2.append((shizi3,daan3))
        k = k+1
    return result2


JIAN_QIAN1000 = Jian_qian1000()


###主程序

if __name__ == "__main__" :


#分别是两种类型题目，所以分开
    result = []
    result2 = []

#随机顺序6种题型，加入到result中来，前面的6种都是全部遍历的，
###，题库只要200题，200题，100题，100题，200题，200题，只要是符合的都在题库当中了
    random.shuffle(JIA_9)
    random.shuffle(JIAN_9)
    random.shuffle(JIA_QIAN)
    random.shuffle(JIAN_QIAN)
    random.shuffle(CHU_JIA_9)
    random.shuffle(CHU_JIAN_9)
    result.extend(JIA_9[0:199])
    result.extend(JIAN_9[0:199])
    result.extend(JIA_QIAN[0:99])
    result.extend(JIAN_QIAN[0:99])
    result.extend(CHU_JIA_9[0:199])
    result.extend(CHU_JIAN_9[0:199])

#各产生1000道题目，然后进行随机顺序加入题库，取前500题
    random.shuffle(JIA_QIAN1000)
    random.shuffle(JIAN_QIAN1000)
    result2.extend(JIA_QIAN1000[0:499])
    result2.extend(JIAN_QIAN1000[0:499])

###再次随机顺序
    random.shuffle(result)
    random.shuffle(result2)  

###生成word的代码
    doc_yy = Document()
    styles = doc_yy.styles
###增加新的模板，字体等
    some1 = styles.add_style('some1',WD_STYLE_TYPE.PARAGRAPH)
    some1.font.size=Pt(14)

####取10题，为一个页面，取result总长度的少10
    i = 0 
    while i < len(result)-10:
        doc_yy.add_paragraph(result[i][0]+' '*50+result2[i][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+1][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+2][0]+' '*50+result2[i+1][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+3][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+4][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+5][0]+' '*50+result2[i+2][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+6][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+7][0],style=some1)
        doc_yy.add_paragraph(text='\r')   
        doc_yy.add_paragraph(result[i+8][0]+' '*50+result2[i+3][0],style=some1)
        doc_yy.add_paragraph(text='\r')
        doc_yy.add_paragraph(result[i+9][0],style=some1)

        i = i + 10

    doc_yy.save('jieguo.docx')


